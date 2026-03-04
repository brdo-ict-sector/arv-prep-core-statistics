import os
import sqlite3
import requests
import json
from typing import List, Optional
from fastapi import FastAPI, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
import tempfile

load_dotenv()

# Configure OpenRouter API
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_MODEL = "google/gemini-3.1-pro-preview"

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DB_PATH = os.path.join(os.path.dirname(__file__), "../data/statistics.db")

def get_db_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

class GenerateRequest(BaseModel):
    kveds: List[str]
    start_year: int
    end_year: int

@app.get("/")
async def root():
    return {"message": "Welcome to State Statistics to ARV API"}

@app.get("/kveds")
async def get_kveds():
    conn = get_db_connection()
    try:
        query = "SELECT DISTINCT kved, kved_name, kved_full FROM active_enterprises ORDER BY kved"
        rows = conn.execute(query).fetchall()
        return [{"kved": row["kved"], "name": row["kved_name"], "full": row["kved_full"]} for row in rows]
    finally:
        conn.close()

@app.post("/generate-docx")
async def generate_docx(request: GenerateRequest):
    conn = get_db_connection()
    try:
        doc = Document()
        doc.add_heading('Основна статистика', level=2)
        
        # New Heading 3 and Source Citation
        doc.add_heading('Кількість підприємств галузі', level=3)
        doc.add_paragraph('За даними Державної служби статистики, "Структурні зміни в економіці України та її регіонів: Показники діяльності підприємств (2012-2024)".')
        
        all_years = [f"year_{y}" for y in range(request.start_year, request.end_year + 1)]
        year_cols = ", ".join(all_years)

        for kved_code in request.kveds:
            # 1. Fetch data for each KVED
            query = f"SELECT business_size, kved_full, {year_cols} FROM active_enterprises WHERE kved = ?"
            rows = conn.execute(query, (kved_code,)).fetchall()
            
            if not rows:
                continue

            kved_full = rows[0]["kved_full"]
            
            # 2. Heading 4 for KVED
            doc.add_heading(kved_full, level=4)
            
            # 3. Add Table
            table_title = f"Кількість підприємств за КВЕД та розміром - за період {request.start_year}-{request.end_year}"
            doc.add_paragraph(table_title)
            
            table = doc.add_table(rows=1, cols=len(all_years) + 1)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Розмір підприємства'
            for i, year in enumerate(range(request.start_year, request.end_year + 1)):
                hdr_cells[i+1].text = str(year)
                
            data_summary = ""
            for row_data in rows:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row_data['business_size'])
                years_info = []
                for i, year in enumerate(range(request.start_year, request.end_year + 1)):
                    val = str(row_data[f'year_{year}'])
                    row_cells[i+1].text = val
                    years_info.append(f"{year}: {val}")
                data_summary += f"Розмір {row_data['business_size']}: {', '.join(years_info)}\n"

            # 4. Generate Strict LLM Description
            prompt = f"""
            Ти — аналітик даних. Напиши сухий, професійний опис динаміки кількості підприємств для звіту.
            КВЕД: {kved_full}
            Період: {request.start_year}-{request.end_year}
            Дані:
            {data_summary}
            
            ВИМОГИ:
            - Тільки фактичний опис тенденцій (зростання, спад, стабільність).
            - БЕЗ вступних фраз (наприклад, "Ось аналіз...", "Згідно з даними...").
            - БЕЗ висновків про регуляторний вплив.
            - БЕЗ патетики.
            - Мова: українська.
            - Максимум 2 абзаци.
            """
            
            description = ""
            if OPENROUTER_API_KEY:
                try:
                    response = requests.post(
                        url="https://openrouter.ai/api/v1/chat/completions",
                        headers={
                            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                            "Content-Type": "application/json",
                        },
                        data=json.dumps({
                            "model": OPENROUTER_MODEL,
                            "messages": [
                                {"role": "system", "content": "Ти — сухий аналітик даних. Пиши лише суть без вступів та підсумків."},
                                {"role": "user", "content": prompt}
                            ]
                        })
                    )
                    if response.status_code == 200:
                        description = response.json()['choices'][0]['message']['content'].strip()
                    else:
                        description = f"[Помилка API: {response.status_code}]"
                except:
                    description = "[Помилка генерації опису]"

            if description:
                doc.add_paragraph(description)
            
            doc.add_paragraph() # Spacer between KVED sections

        if not doc.paragraphs and not doc.tables:
            raise HTTPException(status_code=404, detail="No data found for selected KVEDs")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            return FileResponse(path=tmp.name, filename=f"ARV_Statistics_Multiple.docx", media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
